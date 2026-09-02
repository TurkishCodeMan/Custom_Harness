#!/usr/bin/env python3
"""
High-Speed Python Document Parsing & Chunking Engine
Interacts with Node.js via stdin/stdout line-delimited JSON messages.
"""

import sys
import json
import os
import traceback

def process_single_page_ocr(page_tuple):
    page_num, pix_bytes, ocr_url = page_tuple
    import requests
    import base64
    import time

    b64_img = base64.b64encode(pix_bytes).decode("utf-8")
    payload = {
        "model": "zai-org/GLM-OCR",
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64_img}"}},
                    {"type": "text", "text": "Text Recognition:"}
                ]
            }
        ],
        "max_tokens": 2048
    }

    # Up to 3 attempts with 300s timeout to survive GPU queuing
    for attempt in range(3):
        try:
            res = requests.post(f"{ocr_url}/v1/chat/completions", json=payload, timeout=300)
            if res.status_code == 200:
                data = res.json()
                text = data.get("choices", [{}])[0].get("message", {}).get("content", "").replace('\x00', '').strip()
                if text:
                    return (page_num, f"--- [Sayfa {page_num}] ---\n{text}")

            elif res.status_code == 503:
                time.sleep(2)
        except Exception as e:
            if attempt == 2:
                sys.stderr.write(f"[PythonEngine:OCR] Page {page_num} OCR final error after 3 attempts: {e}\n")
            time.sleep(1)

    return (page_num, "")

def parse_pdf_ocr(file_path):
    try:
        import fitz
        from concurrent.futures import ThreadPoolExecutor
        raw_urls = os.getenv("VLLM_VISION_URL", "http://localhost:8010").split(",")
        ocr_urls = [u.strip().replace("/process", "").replace("/v1", "").rstrip("/") for u in raw_urls if u.strip()]
        if not ocr_urls:
            ocr_urls = ["http://localhost:8010"]

        doc = fitz.open(file_path)
        total_pages = len(doc)
        sys.stderr.write(f"[PythonEngine:OCR] Starting 8-Worker Parallel Robust OCR across {total_pages} pages ({len(ocr_urls)} GPU OCR Havuzu) for: {os.path.basename(file_path)}...\n")

        page_tasks = []
        for i in range(total_pages):
            page = doc.load_page(i)
            # Render page at 1.4x scale for optimal GLM-OCR speed & memory efficiency
            pix = page.get_pixmap(matrix=fitz.Matrix(1.4, 1.4))
            pix_bytes = pix.tobytes(output="jpeg", jpg_quality=85)
            # Round-robin load balancing across all OCR GPUs
            target_url = ocr_urls[i % len(ocr_urls)]
            page_tasks.append((i + 1, pix_bytes, target_url))
        doc.close()

        results = {}
        # Concurrently process 8 pages in parallel (optimized for GPU memory & seq limits)
        with ThreadPoolExecutor(max_workers=8) as executor:
            for page_num, text in executor.map(process_single_page_ocr, page_tasks):
                if text:
                    results[page_num] = text

        # Sort pages in original order
        sorted_pages = [results[i] for i in range(1, total_pages + 1) if i in results]
        return "\n\n".join(sorted_pages)
    except Exception as e:
        sys.stderr.write(f"[PythonEngine:OCR] Parallel OCR failed for {file_path}: {e}\n")
        return ""


def parse_pdf(file_path):
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        pages = []
        full_text = []
        total_chars = 0
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            total_chars += len(text.strip())
            pages.append({
                "page": page_num + 1,
                "text": text
            })
            if text.strip():
                full_text.append(f"--- [Sayfa {page_num + 1}] ---\n{text}")
        page_count = len(doc)
        doc.close()

        extracted_content = "\n\n".join(full_text)

        # Scanned PDF Detection: If average chars per page is less than 30, it's a scanned/photocopied image PDF!
        if page_count > 0 and (total_chars / page_count) < 30:
            sys.stderr.write(f"[PythonEngine] Scanned PDF detected ({total_chars} chars across {page_count} pages): '{os.path.basename(file_path)}'. Routing to 8010 GLM-OCR microservice...\n")
            ocr_text = parse_pdf_ocr(file_path)
            if ocr_text:
                extracted_content = ocr_text


        return {
            "success": True,
            "type": "pdf",
            "pageCount": page_count,
            "content": extracted_content,
            "pages": pages
        }
    except ImportError:
        return {"success": False, "error": "PyMuPDF (fitz) is not installed"}
    except Exception as e:
        return {"success": False, "error": f"PDF parse error: {str(e)}"}


def parse_docx(file_path):
    try:
        import docx
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract table contents
        tables_text = []
        for t in doc.tables:
            for row in t.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    tables_text.append(row_text)
                    
        content = "\n\n".join(paragraphs)
        if tables_text:
            content += "\n\n--- TABLOLAR ---\n" + "\n".join(tables_text)
            
        return {
            "success": True,
            "type": "docx",
            "paragraphCount": len(paragraphs),
            "content": content
        }
    except ImportError:
        return {"success": False, "error": "python-docx is not installed"}
    except Exception as e:
        return {"success": False, "error": f"DOCX parse error: {str(e)}"}

def parse_excel(file_path):
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, data_only=True, read_only=True)
        sheets_content = []
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_rows = []
            for row in ws.iter_rows(values_only=True):
                cleaned_row = [str(val).strip() for val in row if val is not None and str(val).strip() != ""]
                if cleaned_row:
                    sheet_rows.append(" | ".join(cleaned_row))
            if sheet_rows:
                sheets_content.append(f"--- [Sayfa: {sheet_name}] ---\n" + "\n".join(sheet_rows))
                
        wb.close()
        return {
            "success": True,
            "type": "excel",
            "sheetCount": len(sheets_content),
            "content": "\n\n".join(sheets_content)
        }
    except ImportError:
        return {"success": False, "error": "openpyxl is not installed"}
    except Exception as e:
        return {"success": False, "error": f"Excel parse error: {str(e)}"}

def parse_text(file_path):
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        return {
            "success": True,
            "type": "text",
            "content": content
        }
    except Exception as e:
        return {"success": False, "error": f"Text read error: {str(e)}"}

def handle_request(req):
    file_path = req.get("filePath")
    if not file_path or not os.path.exists(file_path):
        return {"id": req.get("id"), "success": False, "error": f"File not found: {file_path}"}
    
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        res = parse_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        res = parse_docx(file_path)
    elif ext in [".xlsx", ".xls", ".csv"]:
        res = parse_excel(file_path)
    else:
        res = parse_text(file_path)
        
    res["id"] = req.get("id")
    res["filePath"] = file_path
    return res

def main():
    try:
        # Signal readiness to parent process
        print(json.dumps({"type": "ready", "pid": os.getpid()}), flush=True)
        
        for line in sys.stdin:
            line = line.strip()
            if not line:
                continue
            try:
                req = json.loads(line)
                if req.get("type") == "ping":
                    print(json.dumps({"id": req.get("id"), "type": "pong"}), flush=True)
                    continue
                    
                res = handle_request(req)
                print(json.dumps(res), flush=True)
            except (BrokenPipeError, IOError):
                break
            except Exception as e:
                try:
                    print(json.dumps({
                        "success": False,
                        "error": f"Worker internal error: {str(e)}",
                        "trace": traceback.format_exc()
                    }), flush=True)
                except (BrokenPipeError, IOError):
                    break
    except (BrokenPipeError, IOError, KeyboardInterrupt):
        pass
    finally:
        try:
            sys.stdout.flush()
        except Exception:
            pass
        sys.exit(0)

if __name__ == "__main__":
    main()

